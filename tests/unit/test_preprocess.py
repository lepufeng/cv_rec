"""Document preprocessor tests."""
from __future__ import annotations

import io

import pytest
from PIL import Image

from app.core.exceptions import ValidationError
from app.parsers.preprocess import _extract_pdf_text, detect_format, preprocess
from tests.fixtures.build_pdf import make_minimal_pdf


def _make_jpeg(width: int = 200, height: int = 200) -> bytes:
    img = Image.new("RGB", (width, height), color=(200, 200, 200))
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=80)
    return buf.getvalue()


def test_detect_format_normalizes_jpeg():
    assert detect_format("foo.JPEG") == "jpg"
    assert detect_format("foo.jpg") == "jpg"
    assert detect_format("foo.PDF") == "pdf"


def test_detect_format_rejects_unknown():
    with pytest.raises(ValidationError):
        detect_format("foo.exe")


def test_preprocess_image():
    data = _make_jpeg()
    result = preprocess("photo.jpg", data)
    assert len(result.images) == 1
    assert result.text is None
    # output should be a valid normalized image
    Image.open(io.BytesIO(result.images[0])).verify()


def test_preprocess_image_resizes_large():
    data = _make_jpeg(width=4096, height=3000)
    result = preprocess("big.jpg", data)
    img = Image.open(io.BytesIO(result.images[0]))
    assert max(img.size) <= 2560


def test_preprocess_bad_image_raises():
    with pytest.raises(ValidationError):
        preprocess("bad.png", b"not-an-image")


def test_preprocess_docx_text_extraction():
    from docx import Document

    doc = Document()
    doc.add_paragraph("张三")
    doc.add_paragraph("13800138000")
    doc.add_paragraph("zhangsan@example.com")
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    result = preprocess("resume.docx", data)
    assert result.text and "张三" in result.text
    assert "13800138000" in result.text
    assert result.images == []


def test_extract_pdf_text_uses_embedded_text_layer():
    text = _extract_pdf_text(make_minimal_pdf("Resume Sample"))
    assert text is not None
    assert "Resume Sample" in text
