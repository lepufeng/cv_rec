"""End-to-end resume upload + parse + patch flow."""
from __future__ import annotations

import io

import pytest
from docx import Document

from tests.fixtures.sample_resume import SAMPLE_PARSED_RESUME


def _make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("张三 - 软件工程师")
    doc.add_paragraph("电话：13800138000")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_register_and_upload_and_get(app_client, make_user):
    client, fake = app_client

    user = await make_user("alice")
    headers = {"Authorization": f"Bearer {user['token']}"}

    fake.queue_response(SAMPLE_PARSED_RESUME)
    files = {"file": (
        "alice.docx",
        _make_docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )}
    resp = await client.post("/api/v1/resumes", headers=headers, files=files)
    assert resp.status_code == 201, resp.text
    detail = resp.json()
    resume_id = detail["resume_id"]
    assert detail["status"] == "completed"
    assert detail["data"]["basic_info"]["name"] == "张三"

    resp = await client.get(f"/api/v1/resumes/{resume_id}", headers=headers)
    assert resp.status_code == 200
    assert resp.json()["data"]["basic_info"]["phone"] == "13800138000"

    resp = await client.patch(
        f"/api/v1/resumes/{resume_id}",
        headers=headers,
        json={"patch": {"basic_info": {"name": "张四"}}},
    )
    assert resp.status_code == 200
    assert resp.json()["data"]["basic_info"]["name"] == "张四"
    assert resp.json()["parsed_data_version"] == 2


@pytest.mark.asyncio
async def test_upload_dedup_by_content_hash(app_client, make_user):
    client, fake = app_client
    user = await make_user("bob")
    headers = {"Authorization": f"Bearer {user['token']}"}

    fake.queue_response(SAMPLE_PARSED_RESUME)
    content = _make_docx_bytes()
    files = {"file": ("a.docx", content, "application/octet-stream")}

    r1 = await client.post("/api/v1/resumes", headers=headers, files=files)
    assert r1.status_code == 201
    first_id = r1.json()["resume_id"]
    initial_calls = len(fake.calls)

    r2 = await client.post(
        "/api/v1/resumes",
        headers=headers,
        files={"file": ("a.docx", content, "application/octet-stream")},
    )
    assert r2.status_code == 201
    assert r2.json()["resume_id"] == first_id
    assert len(fake.calls) == initial_calls


@pytest.mark.asyncio
async def test_delete_resume_cascades(app_client, make_user):
    client, fake = app_client
    user = await make_user("carol")
    headers = {"Authorization": f"Bearer {user['token']}"}

    fake.queue_response(SAMPLE_PARSED_RESUME)
    files = {"file": ("c.docx", _make_docx_bytes(), "application/octet-stream")}
    rid = (await client.post("/api/v1/resumes", headers=headers, files=files)).json()["resume_id"]

    resp = await client.delete(f"/api/v1/resumes/{rid}", headers=headers)
    assert resp.status_code == 204

    resp = await client.get(f"/api/v1/resumes/{rid}", headers=headers)
    assert resp.status_code == 404
